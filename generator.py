import random
import json

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

def read_words(file_path):
    with open(file_path, 'r') as file:
        words = [line.strip().upper() for line in file.readlines()]
    return words

def calculate_letter_frequencies(words):
    freq = {}
    total_letters = 0

    for word in words:
        for char in word:
            freq[char] = freq.get(char, 0) + 1
            total_letters += 1

    for char in freq:
        freq[char] /= total_letters

    return freq

def fill_empty_spaces(grid, freq):
    for i in range(len(grid)):
        for j in range(len(grid[i])):
            if grid[i][j] == ' ':
                rand_value = random.random()
                cumulative_prob = 0.0
                for char, prob in freq.items():
                    cumulative_prob += prob
                    if rand_value < cumulative_prob:
                        grid[i][j] = char
                        break

def place_word_in_grid(grid, word, direction, max_attempt=400):

    col_dir, row_dir = direction

    max_row = len(grid) - len(word) * row_dir
    max_col = len(grid[0]) - len(word) * col_dir
    
    attempt = 0

    while attempt < max_attempt:
    
        start_row = random.randint(0, max_row)
        start_col = random.randint(0, max_col)

        fits=True
        for i in range(len(word)):
            new_row = start_row + i * row_dir
            new_col = start_col + i * col_dir

            if new_row < 0 or new_row >= len(grid) or new_col < 0 or new_col >= len(grid[0]):
                fits = False
            elif grid[new_row][new_col] != ' ' and grid[new_row][new_col] != word[i]:
                fits = False
                
        if fits:
            for i in range(len(word)):
                new_row = start_row + i * row_dir
                new_col = start_col + i * col_dir
                grid[new_row][new_col] = word[i]
            return True
        else:
            attempt += 1
    
    return False
        
def generate_word_search():
    grid = [[' ' for _ in range(16)] for _ in range(16)]
    
    word_list = read_words('word_list.txt')
    selected = random.sample([w for w in word_list if len(w) == 4], 2) + \
            random.sample([w for w in word_list if len(w) == 5], 3) + \
            random.sample([w for w in word_list if len(w) == 6], 2) + \
            random.sample([w for w in word_list if len(w) == 8], 2) + \
            random.sample([w for w in word_list if len(w) == 10], 1)
    directions = [(1, 0), (1, 0), (0, 1), (0, 1), (-1, 0), (0, -1), (1, 1), (1, -1), (-1, 1), (-1, -1)]
    random.shuffle(directions)

    for word, direction in zip(selected, directions):
        if not place_word_in_grid(grid, word, direction):
            return None, None
        
    letter_freq = calculate_letter_frequencies(selected)
    fill_empty_spaces(grid, letter_freq)
    
    return grid, selected

def set_narrow_margins(document):
    sections = document.sections
    for section in sections:
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)

def save_puzzles_to_word(num_puzzles=10, output_file='word_search_puzzles.docx', json_file='puzzle_words.json'):
    document = Document()
    set_narrow_margins(document)

    puzzles_data = {}

    for i in range(num_puzzles):
        grid, selected = generate_word_search()
        if grid is None:
            print(f"Puzzle {i + 1} could not be generated.")
            continue

        puzzles_data[f"puzzle {i+1}"] = selected
        document.add_heading(f"Puzzle {i + 1}", level=1)
        table = document.add_table(rows=16, cols=16)

        for row_idx, row in enumerate(grid):
            for col_idx, letter in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = letter.upper()
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name = "맑은 고딕"

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 0.75

        # document.add_paragraph("\nWords:").alignment = WD_ALIGN_PARAGRAPH.CENTER
        # words_paragraph = document.add_paragraph(', '.join(selected))
        # words_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # for run in words_paragraph.runs:
        #     run.font.name = "맑은 고딕"

        document.add_page_break()

    document.save(output_file)
    print(f"{num_puzzles} puzzles saved to {output_file}")

    with open(json_file, 'w') as json_out:
        json.dump(puzzles_data, json_out, indent=4)
    print(f"Puzzle words saved to {json_file}")

save_puzzles_to_word()